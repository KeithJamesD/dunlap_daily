#!/usr/bin/env python3
"""
Test script for Dunlap Daily Newsletter Generator
"""

import sys
import os
from pathlib import Path
import datetime
from docx import Document

# Add the current directory to Python path
sys.path.insert(0, str(Path(__file__).parent))

from dunlap_daily import DunlapDailyGenerator

def create_sample_document():
    """Create a sample .docx file for testing."""
    doc = Document()
    
    # Add title
    title = doc.add_paragraph(f"Daily Update - {datetime.date.today().strftime('%B %d, %Y')}")
    title.runs[0].bold = True
    
    # Add highlights section
    doc.add_paragraph("HIGHLIGHTS")
    doc.add_paragraph("• Completed newsletter automation system")
    doc.add_paragraph("• Set up OneDrive integration")
    doc.add_paragraph("• Configured GitHub Pages deployment")
    
    # Add news section
    doc.add_paragraph("NEWS")
    doc.add_paragraph("The Dunlap Daily newsletter system is now fully operational! This automated system fetches content from OneDrive documents and generates beautiful RSS feeds and web pages.")
    
    # Add announcements section
    doc.add_paragraph("ANNOUNCEMENTS")
    doc.add_paragraph("Starting tomorrow, the newsletter will be automatically generated and published every morning at 9 AM. Make sure to update the daily content document in OneDrive to see your updates reflected on the website.")
    
    # Save the document
    doc_path = Path(__file__).parent / "daily_content.docx"
    doc.save(str(doc_path))
    print(f"Sample document created: {doc_path}")
    
    return doc_path

def test_generator():
    """Test the newsletter generator."""
    print("Testing Dunlap Daily Newsletter Generator")
    print("="*50)
    
    try:
        # Create sample document
        create_sample_document()
        
        # Initialize generator
        generator = DunlapDailyGenerator()
        
        # Run the generator
        generator.run(force_update=True)
        
        # Check output
        output_dir = Path(__file__).parent / "docs"
        
        files_to_check = [
            output_dir / "index.html",
            output_dir / "archive.html", 
            output_dir / "feed.xml",
            output_dir / "style.css",
            output_dir / "entries.json"
        ]
        
        print("\\nChecking generated files:")
        for file_path in files_to_check:
            if file_path.exists():
                print(f"✅ {file_path.name}")
            else:
                print(f"❌ {file_path.name}")
        
        # Check if today's archive file was created
        today_archive = output_dir / "archive" / f"{datetime.date.today().strftime('%Y-%m-%d')}.html"
        if today_archive.exists():
            print(f"✅ {today_archive.name}")
        else:
            print(f"❌ {today_archive.name}")
            
        print(f"\\nOutput directory: {output_dir}")
        print("Test completed!")
        
    except Exception as e:
        print(f"❌ Test failed: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_generator()