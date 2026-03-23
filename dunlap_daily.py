#!/usr/bin/env python3
"""
Dunlap Daily RSS Newsletter Generator

This script fetches daily content from a OneDrive .doc file,
generates an RSS feed, and creates HTML pages for GitHub Pages hosting.
"""

import os
import sys
import json
import datetime
import re
from pathlib import Path
from typing import Dict, List, Optional
import requests
import zipfile
import xml.etree.ElementTree as ET
from xml.dom import minidom
import logging
import argparse
import pytz
import webbrowser
import time
import threading
from http.server import HTTPServer, SimpleHTTPRequestHandler
import socketserver

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Installing...")
    os.system("pip install python-docx")
    from docx import Document

try:
    import feedgen.feed
except ImportError:
    print("feedgen not installed. Installing...")
    os.system("pip install feedgen")
    import feedgen.feed

from feedgen.feed import FeedGenerator

# Try to import email functionality
try:
    from newsletter_email import NewsletterEmailer
    EMAIL_AVAILABLE = True
except ImportError:
    EMAIL_AVAILABLE = False
    print("Email functionality not available. Install required packages if needed.")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DunlapDailyGenerator:
    """Main class for generating the Dunlap Daily RSS newsletter."""
    
    def __init__(self, config_path: str = "config.json"):
        self.config = self.load_config(config_path)
        self.base_dir = Path(__file__).parent
        self.output_dir = self.base_dir / "docs"  # GitHub Pages directory
        self.rss_file = self.output_dir / "feed.xml"
        self.archive_dir = self.output_dir / "archive"
        
        # Create necessary directories
        self.output_dir.mkdir(exist_ok=True)
        self.archive_dir.mkdir(exist_ok=True)
        
    def load_config(self, config_path: str) -> Dict:
        """Load configuration from JSON file."""
        try:
            with open(config_path, 'r') as f:
                return json.load(f)
        except FileNotFoundError:
            logger.warning(f"Config file {config_path} not found. Using defaults.")
            return self.get_default_config()
    
    def get_default_config(self) -> Dict:
        """Return default configuration."""
        return {
            "newsletter": {
                "title": "Dunlap Daily",
                "description": "Daily insights and updates from Dunlap",
                "author": "Dunlap Team",
                "link": "http://localhost:8000",
                "language": "en-US"
            },
            "onedrive": {
                "client_id": "",
                "client_secret": "",
                "tenant_id": "",
                "file_path": "/path/to/your/daily-content.docx"
            },
            "github": {
                "repository": "your-username/dunlap_daily",
                "pages_url": "https://your-username.github.io/dunlap_daily"
            },
            "local_server": {
                "port": 8000,
                "auto_open": true,
                "watch_files": true
            }
        }
    
    def fetch_onedrive_content(self) -> Optional[str]:
        """Fetch content from OneDrive .doc file."""
        try:
            # Try OneDrive integration first
            if self._has_onedrive_config():
                content = self._fetch_from_onedrive()
                if content:
                    return content
            
            # Fallback to local files
            local_doc_path = self.base_dir / "daily_content.docx"
            local_txt_path = self.base_dir / "daily_content.txt"
            
            # Try .docx file first
            if local_doc_path.exists():
                try:
                    # Read the Word document
                    doc = Document(local_doc_path)
                    content = []
                    
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            content.append(paragraph.text.strip())
                    
                    return "\n\n".join(content)
                except Exception as e:
                    logger.warning(f"Error reading .docx file: {e}")
            
            # Try .txt file as fallback
            if local_txt_path.exists():
                try:
                    # Try UTF-8 first
                    with open(local_txt_path, 'r', encoding='utf-8') as f:
                        return f.read()
                except UnicodeDecodeError as e:
                    logger.warning(f"UTF-8 decode error in .txt file: {e}")
                    # Try other common encodings
                    for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                        try:
                            with open(local_txt_path, 'r', encoding=encoding) as f:
                                content = f.read()
                                logger.info(f"Successfully read file using {encoding} encoding")
                                return content
                        except (UnicodeDecodeError, Exception):
                            continue
                    logger.error(f"Could not read .txt file with any supported encoding")
                except Exception as e:
                    logger.warning(f"Error reading .txt file: {e}")
            
            # Check if the .docx file is actually a text file
            if local_doc_path.exists():
                try:
                    with open(local_doc_path, 'r', encoding='utf-8') as f:
                        return f.read()
                except UnicodeDecodeError:
                    # It's a binary file, but not a valid .docx
                    logger.warning(f"File {local_doc_path} appears to be corrupted or not a valid format")
            
            logger.warning(f"No content file found. Expected: {local_doc_path} or {local_txt_path}")
            return None
            
        except Exception as e:
            logger.error(f"Error fetching content: {e}")
            return None
    
    def _has_onedrive_config(self) -> bool:
        """Check if OneDrive configuration is available."""
        onedrive_config = self.config.get("onedrive", {})
        required_keys = ["client_id", "client_secret", "tenant_id", "refresh_token"]
        return all(key in onedrive_config and onedrive_config[key] for key in required_keys)
    
    def _fetch_from_onedrive(self) -> Optional[str]:
        """Fetch content directly from OneDrive using API."""
        try:
            from onedrive_helper import OneDriveClient
            
            onedrive_config = self.config["onedrive"]
            client = OneDriveClient(
                onedrive_config["client_id"],
                onedrive_config["client_secret"], 
                onedrive_config["tenant_id"]
            )
            
            # Refresh token to get current access token
            token_response = client.refresh_token(onedrive_config["refresh_token"])
            access_token = token_response["access_token"]
            
            # Update refresh token if provided
            if "refresh_token" in token_response:
                self.config["onedrive"]["refresh_token"] = token_response["refresh_token"]
                self._save_config()
            
            # Download file
            file_path = onedrive_config.get("file_path")
            if not file_path:
                logger.error("OneDrive file path not configured")
                return None
            
            # Replace {date} placeholder with current date in M_DD_YY format
            if "{date}" in file_path:
                current_date = datetime.now().strftime("%m_%d_%y").lstrip("0").replace("_0", "_")
                file_path = file_path.replace("{date}", current_date)
                logger.info(f"Using dynamic file path: {file_path}")
            
            file_content = client.download_file(file_path, access_token)
            
            # Save to local file for backup
            local_doc_path = self.base_dir / "daily_content.docx"
            with open(local_doc_path, 'wb') as f:
                f.write(file_content)
            
            # Parse the downloaded content
            doc = Document(local_doc_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            logger.info("Successfully fetched content from OneDrive")
            return "\n\n".join(content)
            
        except Exception as e:
            logger.warning(f"Failed to fetch from OneDrive: {e}")
            return None
    
    def _save_config(self):
        """Save updated configuration."""
        try:
            config_path = self.base_dir / "config.json"
            with open(config_path, 'w') as f:
                json.dump(self.config, f, indent=2)
        except Exception as e:
            logger.warning(f"Failed to save config: {e}")
    
    def parse_content(self, raw_content: str, scheduled_date: datetime.datetime = None) -> Dict:
        """Parse the raw content into structured newsletter data."""
        if not raw_content:
            return {}
        
        lines = raw_content.split('\n')
        title = ""
        content_sections = []
        current_section = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # First non-empty line is typically the title
            if not title:
                # Remove zero-width no-break space and other problematic characters
                title = line.replace('\ufeff', '').replace('\u200b', '').strip()
                continue
            
            # Look for section headers (lines that might be in all caps or have special formatting)
            if line.isupper() and len(line) > 3:
                if current_section:
                    content_sections.append(current_section)
                current_section = f"<h3>{line}</h3>\n"
            # Check if line contains HTML tags (like images) - pass through directly
            elif line.startswith('<') and line.endswith('>'):
                current_section += f"{line}\n"
            # Check for image syntax: ![alt text](image_path)
            elif line.startswith('![') and '](' in line and line.endswith(')'):
                # Convert markdown-style image to HTML
                alt_start = line.find('![') + 2
                alt_end = line.find('](')
                src_start = alt_end + 2
                src_end = line.find(')', src_start)
                
                alt_text = line[alt_start:alt_end]
                src_path = line[src_start:src_end]
                
                current_section += f'<img src="{src_path}" alt="{alt_text}" style="max-width: 100%; height: auto; margin: 10px 0;">\n'
            else:
                current_section += f"<p>{line}</p>\n"
        
        if current_section:
            content_sections.append(current_section)
        
        # Use scheduled date if provided, otherwise current time
        pub_date = scheduled_date or datetime.datetime.now(pytz.UTC)
        
        return {
            "title": title or f"Daily Update - {pub_date.strftime('%B %d, %Y')}",
            "content": "\n".join(content_sections),
            "date": pub_date,
            "permalink": f"/archive/{pub_date.strftime('%Y-%m-%d')}.html"
        }
    
    def generate_rss_feed(self, entries: List[Dict]) -> None:
        """Generate RSS feed from entries."""
        fg = FeedGenerator()
        
        # Set feed metadata
        fg.title(self.config["newsletter"]["title"])
        fg.description(self.config["newsletter"]["description"])
        fg.link(href=self.config["newsletter"]["link"], rel="alternate")
        fg.link(href=f"{self.config['newsletter']['link']}/feed.xml", rel="self")
        fg.author(name=self.config["newsletter"]["author"])
        fg.language(self.config["newsletter"]["language"])
        fg.lastBuildDate(datetime.datetime.now(pytz.UTC))
        
        # Add entries to feed
        for entry in entries[:10]:  # Limit to 10 most recent entries
            fe = fg.add_entry()
            fe.title(entry["title"])
            fe.description(entry["content"])
            fe.link(href=f"{self.config['newsletter']['link']}{entry['permalink']}")
            fe.pubDate(entry["date"])
            fe.guid(f"{self.config['newsletter']['link']}{entry['permalink']}", permalink=True)
        
        # Write RSS feed
        fg.rss_file(str(self.rss_file))
        logger.info(f"RSS feed generated: {self.rss_file}")
    
    def generate_html_page(self, entry: Dict) -> None:
        """Generate HTML page for a single entry."""
        date_str = entry["date"].strftime("%Y-%m-%d")
        html_file = self.archive_dir / f"{date_str}.html"
        
        # Fix image paths for archive pages (need ../ prefix)
        content = entry['content']
        content = content.replace('src="images/', 'src="../images/')
        
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{entry['title']} - {self.config['newsletter']['title']}</title>
    <link rel="stylesheet" href="../style.css">
</head>
<body>
    <header>
        <h1><a href="../">{self.config['newsletter']['title']}</a></h1>
        <nav>
            <a href="../">Home</a> | 
            <a href="../feed.xml">RSS Feed</a> | 
            <a href="../archive.html">Archive</a>
        </nav>
    </header>
    
    <main>
        <article>
            <h2>{entry['title']}</h2>
            <time datetime="{entry['date'].isoformat()}">{entry['date'].strftime('%B %d, %Y')}</time>
            <div class="content">
                {content}
            </div>
        </article>
    </main>
    
    <footer>
        <p>&copy; {datetime.date.today().year} {self.config['newsletter']['author']}. All rights reserved.</p>
    </footer>
</body>
</html>
"""
        
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info(f"HTML page generated: {html_file}")

    class LocalServerHandler(SimpleHTTPRequestHandler):
        """Custom handler for local server."""
        def __init__(self, *args, docs_dir=None, **kwargs):
            self.docs_dir = docs_dir
            super().__init__(*args, **kwargs)
        
        def translate_path(self, path):
            """Translate URL path to filesystem path."""
            if self.docs_dir:
                # Remove leading slash and resolve relative to docs directory
                path = path.lstrip('/')
                if not path:
                    path = 'index.html'
                full_path = self.docs_dir / path
                return str(full_path)
            return super().translate_path(path)
    
    def start_local_server(self, port: int = 8000, auto_open: bool = True) -> None:
        """Start a local web server to serve the newsletter."""
        try:
            # Change to docs directory to serve files
            os.chdir(self.output_dir)
            
            # Create a handler that serves from docs directory
            def handler(*args, **kwargs):
                return self.LocalServerHandler(*args, docs_dir=self.output_dir, **kwargs)
            
            # Try to find an available port
            original_port = port
            max_attempts = 10
            
            for attempt in range(max_attempts):
                try:
                    with socketserver.TCPServer(("", port), handler) as httpd:
                        server_url = f"http://localhost:{port}"
                        
                        if port != original_port:
                            logger.info(f"Port {original_port} was busy, using port {port} instead")
                        
                        logger.info(f"Serving newsletter at {server_url}")
                        
                        if auto_open:
                            # Open browser after a short delay
                            threading.Timer(1, lambda: webbrowser.open(server_url)).start()
                        
                        print(f"\\n🌐 Newsletter server running at: {server_url}")
                        print("📧 Press Ctrl+C to stop the server")
                        print("🔄 Server will automatically reload when files change\\n")
                        
                        try:
                            httpd.serve_forever()
                        except KeyboardInterrupt:
                            print("\\n✅ Server stopped.")
                        return
                        
                except OSError as e:
                    if e.errno == 10048:  # Port in use
                        port += 1
                        if attempt < max_attempts - 1:
                            continue
                        else:
                            raise Exception(f"Could not find an available port after trying {original_port}-{port}")
                    else:
                        raise
                    
        except Exception as e:
            logger.error(f"Error starting local server: {e}")
            if "10048" in str(e) or "address already in use" in str(e).lower():
                print(f"\\n❌ Port {original_port} is already in use!")
                print("🔧 Solutions:")
                print(f"   1. Try a different port: python dunlap_daily.py --serve --port {original_port + 1}")
                print("   2. Close other applications using the port")
                print("   3. Kill existing server processes")
                print("\\n💡 The server will automatically try nearby ports next time")
        finally:
            # Change back to original directory
            os.chdir(self.base_dir)
    
    def watch_and_regenerate(self) -> None:
        """Watch for file changes and regenerate content."""
        import time
        last_modified_docx = 0
        last_modified_txt = 0
        content_file_docx = self.base_dir / "daily_content.docx"
        content_file_txt = self.base_dir / "daily_content.txt"
        
        logger.info("👀 Watching for content changes...")
        
        while True:
            try:
                regenerate_needed = False
                
                # Check .docx file
                if content_file_docx.exists():
                    current_modified = content_file_docx.stat().st_mtime
                    if current_modified > last_modified_docx:
                        if last_modified_docx > 0:  # Skip initial run
                            regenerate_needed = True
                        last_modified_docx = current_modified
                
                # Check .txt file
                if content_file_txt.exists():
                    current_modified = content_file_txt.stat().st_mtime
                    if current_modified > last_modified_txt:
                        if last_modified_txt > 0:  # Skip initial run
                            regenerate_needed = True
                        last_modified_txt = current_modified
                
                if regenerate_needed:
                    logger.info("📝 Content file changed - regenerating...")
                    self.run(force_update=True)
                    logger.info("✅ Newsletter regenerated")
                
                time.sleep(2)  # Check every 2 seconds
                
            except KeyboardInterrupt:
                break
            except Exception as e:
                logger.warning(f"Error watching files: {e}")
                time.sleep(5)
    
    def serve_local(self, port: int = None, auto_open: bool = None, watch: bool = None) -> None:
        """Serve the newsletter locally with optional file watching."""
        # Use config defaults if not specified
        local_config = self.config.get("local_server", {})
        port = port or local_config.get("port", 8000)
        auto_open = auto_open if auto_open is not None else local_config.get("auto_open", True)
        watch = watch if watch is not None else local_config.get("watch_files", True)
        
        # Generate initial content
        logger.info("🚀 Generating newsletter content...")
        self.run()
        
        # Start file watcher in a separate thread if enabled
        if watch:
            watch_thread = threading.Thread(target=self.watch_and_regenerate, daemon=True)
            watch_thread.start()
        
        # Start the local server
        self.start_local_server(port, auto_open)
    
    def load_existing_entries(self) -> List[Dict]:
        """Load existing entries from archive files."""
        entries = []
        entries_file = self.output_dir / "entries.json"
        
        if entries_file.exists():
            try:
                with open(entries_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    for item in data:
                        item["date"] = datetime.datetime.fromisoformat(item["date"])
                        entries.append(item)
            except Exception as e:
                logger.error(f"Error loading existing entries: {e}")
        
        return entries
    
    def save_entries(self, entries: List[Dict]) -> None:
        """Save entries to JSON file."""
        entries_file = self.output_dir / "entries.json"
        
        # Convert datetime objects to strings for JSON serialization
        serializable_entries = []
        for entry in entries:
            entry_copy = entry.copy()
            entry_copy["date"] = entry["date"].isoformat()
            serializable_entries.append(entry_copy)
        
        with open(entries_file, 'w', encoding='utf-8') as f:
            json.dump(serializable_entries, f, indent=2, ensure_ascii=False)
    
    def generate_index_page(self, entries: List[Dict]) -> None:
        """Generate the main index page."""
        index_file = self.output_dir / "index.html"
        
        # Generate recent entries HTML
        recent_entries_html = ""
        for entry in entries[:5]:  # Show 5 most recent
            date_str = entry["date"].strftime("%B %d, %Y")
            recent_entries_html += f"""
            <article class="entry-full">
                <h3>{entry['title']}</h3>
                <time>{date_str}</time>
                <div class="content">
                    {entry['content']}
                </div>
            </article>
            """
        
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.config['newsletter']['title']}</title>
    <meta name="description" content="{self.config['newsletter']['description']}">
    <link rel="stylesheet" href="style.css">
    <link rel="alternate" type="application/rss+xml" title="RSS Feed" href="feed.xml">
</head>
<body>
    <header>
        <h1>{self.config['newsletter']['title']}</h1>
        <p class="tagline">{self.config['newsletter']['description']}</p>
        <nav>
            <a href="feed.xml">RSS Feed</a> | 
            <a href="archive.html">Archive</a>
        </nav>
    </header>
    
    <main>
        <section class="recent-entries">
            <h2>Recent Updates</h2>
            {recent_entries_html}
        </section>
    </main>
    
    <footer>
        <p>&copy; {datetime.date.today().year} {self.config['newsletter']['author']}. All rights reserved.</p>
        <p><a href="feed.xml">Subscribe to RSS</a></p>
    </footer>
</body>
</html>
"""
        
        with open(index_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info(f"Index page generated: {index_file}")
    
    def generate_archive_page(self, entries: List[Dict]) -> None:
        """Generate the archive page."""
        archive_file = self.output_dir / "archive.html"
        
        # Generate archive entries HTML
        archive_entries_html = ""
        for entry in entries:
            date_str = entry["date"].strftime("%B %d, %Y")
            archive_entries_html += f"""
            <li>
                <time>{date_str}</time>
                <a href="{entry['permalink']}">{entry['title']}</a>
            </li>
            """
        
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Archive - {self.config['newsletter']['title']}</title>
    <link rel="stylesheet" href="style.css">
</head>
<body>
    <header>
        <h1><a href="./">{self.config['newsletter']['title']}</a></h1>
        <nav>
            <a href="./">Home</a> | 
            <a href="feed.xml">RSS Feed</a> | 
            <a href="archive.html">Archive</a>
        </nav>
    </header>
    
    <main>
        <h2>Archive</h2>
        <ul class="archive-list">
            {archive_entries_html}
        </ul>
    </main>
    
    <footer>
        <p>&copy; {datetime.date.today().year} {self.config['newsletter']['author']}. All rights reserved.</p>
    </footer>
</body>
</html>
"""
        
        with open(archive_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info(f"Archive page generated: {archive_file}")
    
    def run(self, force_update: bool = False, scheduled_date: datetime.datetime = None) -> None:
        """Main execution method."""
        logger.info("Starting Dunlap Daily generation...")
        
        # Load existing entries
        entries = self.load_existing_entries()
        
        # Determine target date for new entry
        target_date = scheduled_date.date() if scheduled_date else datetime.date.today()
        
        # Check if we need to add a new entry for target date
        has_target_entry = any(
            entry["date"].date() == target_date for entry in entries
        )
        
        if not has_target_entry or force_update:
            # Fetch new content
            raw_content = self.fetch_onedrive_content()
            
            if raw_content:
                # Parse content with scheduled date
                new_entry = self.parse_content(raw_content, scheduled_date)
                
                # Add to entries list
                entries.append(new_entry)
                
                # Sort entries by date (newest first)
                entries.sort(key=lambda x: x["date"], reverse=True)
                
                # Generate HTML page for new entry
                self.generate_html_page(new_entry)
                
                scheduled_info = f" (scheduled for {scheduled_date.strftime('%B %d, %Y at %I:%M %p')})" if scheduled_date else ""
                logger.info(f"New entry added: {new_entry['title']}{scheduled_info}")
            else:
                logger.warning("No new content found")
        else:
            target_info = f"for {target_date.strftime('%B %d, %Y')}" if scheduled_date else "for today"
            logger.info(f"Entry {target_info} already exists")
        
        # Generate RSS feed
        self.generate_rss_feed(entries)
        
        # Generate pages
        self.generate_index_page(entries)
        self.generate_archive_page(entries)
        
        # Save entries
        self.save_entries(entries)
        
        # Send email newsletter if enabled and new content was added
        if EMAIL_AVAILABLE and (not has_target_entry or force_update) and raw_content:
            self.send_email_newsletter(new_entry)
        
        logger.info("Dunlap Daily generation complete!")
    
    def send_email_newsletter(self, entry: Dict) -> None:
        """Send newsletter via email if configured."""
        try:
            emailer = NewsletterEmailer()
            success = emailer.send_newsletter(entry)
            if success:
                logger.info("Newsletter sent via email successfully")
            else:
                logger.warning("Email sending failed or disabled")
        except Exception as e:
            logger.error(f"Error sending email newsletter: {e}")

def parse_schedule_string(schedule_str: str) -> Optional[datetime.datetime]:
    """Parse schedule string into datetime object."""
    schedule_str = schedule_str.lower().strip()
    
    try:
        # Handle "tomorrow noon" specifically
        if "tomorrow noon" in schedule_str:
            tomorrow = datetime.date.today() + datetime.timedelta(days=1)
            return datetime.datetime.combine(tomorrow, datetime.time(12, 0)).replace(tzinfo=pytz.UTC)
        
        # Handle "tomorrow" with custom time
        if schedule_str.startswith("tomorrow"):
            tomorrow = datetime.date.today() + datetime.timedelta(days=1)
            # Extract time part if present
            if " " in schedule_str:
                time_part = schedule_str.split(" ", 1)[1]
                if "noon" in time_part:
                    return datetime.datetime.combine(tomorrow, datetime.time(12, 0)).replace(tzinfo=pytz.UTC)
                elif "midnight" in time_part:
                    return datetime.datetime.combine(tomorrow, datetime.time(0, 0)).replace(tzinfo=pytz.UTC)
                else:
                    # Try to parse time like "3pm", "14:30", etc.
                    try:
                        if "pm" in time_part or "am" in time_part:
                            time_obj = datetime.datetime.strptime(time_part.replace("pm", " PM").replace("am", " AM"), "%I %p").time()
                        else:
                            time_obj = datetime.datetime.strptime(time_part, "%H:%M").time()
                        return datetime.datetime.combine(tomorrow, time_obj).replace(tzinfo=pytz.UTC)
                    except ValueError:
                        pass
            # Default to noon if just "tomorrow"
            return datetime.datetime.combine(tomorrow, datetime.time(12, 0)).replace(tzinfo=pytz.UTC)
        
        # Handle full date-time format: "YYYY-MM-DD HH:MM"
        if "-" in schedule_str and ":" in schedule_str:
            return datetime.datetime.strptime(schedule_str, "%Y-%m-%d %H:%M").replace(tzinfo=pytz.UTC)
        
        # Handle date only: "YYYY-MM-DD" (default to noon)
        if "-" in schedule_str:
            date_obj = datetime.datetime.strptime(schedule_str, "%Y-%m-%d").date()
            return datetime.datetime.combine(date_obj, datetime.time(12, 0)).replace(tzinfo=pytz.UTC)
        
    except ValueError:
        pass
    
    return None

def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(description="Generate Dunlap Daily RSS Newsletter")
    parser.add_argument(
        "--force", 
        action="store_true", 
        help="Force update even if today's entry exists"
    )
    parser.add_argument(
        "--config", 
        default="config.json", 
        help="Path to configuration file"
    )
    parser.add_argument(
        "--serve", 
        action="store_true", 
        help="Start local server to serve the newsletter"
    )
    parser.add_argument(
        "--port", 
        type=int, 
        default=None,
        help="Port for local server (default: 8000)"
    )
    parser.add_argument(
        "--no-watch",
        action="store_true",
        help="Disable file watching when serving locally"
    )
    parser.add_argument(
        "--no-open",
        action="store_true", 
        help="Don't automatically open browser when serving locally"
    )
    parser.add_argument(
        "--schedule",
        type=str,
        help="Schedule post for specific date/time (format: 'YYYY-MM-DD HH:MM' or 'tomorrow noon')"
    )
    parser.add_argument(
        "--email-add",
        type=str,
        help="Add email subscriber"
    )
    parser.add_argument(
        "--email-remove", 
        type=str,
        help="Remove email subscriber"
    )
    parser.add_argument(
        "--email-list",
        action="store_true",
        help="List active email subscribers"
    )
    parser.add_argument(
        "--email-test",
        type=str, 
        help="Send test email to address"
    )
    
    args = parser.parse_args()
    
    try:
        generator = DunlapDailyGenerator(args.config)
        
        # Parse scheduled date if provided
        scheduled_date = None
        if args.schedule:
            scheduled_date = parse_schedule_string(args.schedule)
            if not scheduled_date:
                logger.error(f"Invalid schedule format: {args.schedule}")
                logger.error("Examples: 'tomorrow noon', '2026-03-19 12:00', 'tomorrow 3pm'")
                sys.exit(1)
        
        # Handle email commands
        if args.email_add or args.email_remove or args.email_list or args.email_test:
            if not EMAIL_AVAILABLE:
                logger.error("Email functionality not available. Please check newsletter_email.py")
                sys.exit(1)
            
            emailer = NewsletterEmailer()
            
            if args.email_add:
                success = emailer.add_subscriber(args.email_add)
                print(f"{'Added' if success else 'Failed to add'} subscriber: {args.email_add}")
                
            elif args.email_remove:
                success = emailer.remove_subscriber(args.email_remove)
                print(f"{'Removed' if success else 'Failed to remove'} subscriber: {args.email_remove}")
                
            elif args.email_list:
                subscribers = emailer.get_active_subscribers()
                print(f"Active subscribers ({len(subscribers)}):")
                for email in subscribers:
                    print(f"  - {email}")
                    
            elif args.email_test:
                # Create a test entry
                test_entry = {
                    "title": "Test Newsletter - Dunlap Daily",
                    "content": "<p>This is a test email from your Dunlap Daily newsletter system.</p><p>If you're receiving this, email functionality is working correctly!</p>",
                    "date": datetime.datetime.now(pytz.UTC)
                }
                success = emailer.send_newsletter(test_entry, test_email=args.email_test)
                print(f"Test email {'sent successfully' if success else 'failed'}")
            
            return
        
        if args.serve:
            # Start local server
            generator.serve_local(
                port=args.port,
                auto_open=not args.no_open,
                watch=not args.no_watch
            )
        else:
            # Standard generation
            generator.run(force_update=args.force, scheduled_date=scheduled_date)
    except Exception as e:
        logger.error(f"Error running generator: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
